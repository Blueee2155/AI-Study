# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125))
    elif level == 2:
        set_font(run, size=14, bold=True, color=(68, 114, 196))
    elif level == 3:
        set_font(run, size=12, bold=True, color=(68, 114, 196))
    return heading


def add_normal_para(doc, text, size=10.5, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_table_simple(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_shading(hdr_cells[i], '4472C4')
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    set_cell_shading_for_para(p, 'F2F2F2')
    return p


def set_cell_shading_for_para(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def generate_doc1():
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('考研助手智能体项目立项与架构说明书')
    set_font(run, size=22, bold=True, color=(31, 73, 125))
    title.paragraph_format.space_after = Pt(12)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'
    info_data = [
        ('项目名称', 'AI 学习助手（考研助手）'),
        ('文档版本', 'V1.0'),
        ('编写日期', '2026-07-03'),
        ('文档状态', '正式发布'),
    ]
    for i, (k, v) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = ''
        p1 = row[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_font(r1, size=10.5, bold=True)
        set_cell_shading(row[0], 'D9E2F3')
        
        row[1].text = ''
        p2 = row[1].paragraphs[0]
        r2 = p2.add_run(v)
        set_font(r2, size=10.5)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '一、项目概述', level=1)
    
    add_heading_styled(doc, '1.1 项目背景', level=2)
    add_normal_para(doc, '考研竞争日益激烈，备考过程漫长且枯燥，学习者普遍面临两个核心问题：')
    
    p = doc.add_paragraph(style='List Number')
    run = p.add_run('学习状态难以维持：长时间自习容易走神、打瞌睡，缺乏外部监督机制')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    p = doc.add_paragraph(style='List Number')
    run = p.add_run('疑难问题即时解答：复习中遇到知识点疑惑时，缺乏专业、及时的辅导资源')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    add_normal_para(doc, '本项目构建一款 Web 端 AI 学习助手，通过视觉感知与大模型推理的智能体架构，实现"学习状态监督 + 考研智能问答"双重功能，帮助考研学生提高学习效率。')
    
    add_heading_styled(doc, '1.2 项目目标', level=2)
    add_table_simple(doc,
        ['目标层级', '描述', '验收指标'],
        [
            ['核心目标1', '实时学习状态监督', '摄像头检测人脸位置与状态，识别"专注/分心/开小差/离开"四种状态，准确率 ≥ 85%'],
            ['核心目标2', '考研智能问答', '集成 DeepSeek 大模型 + RAG 知识库，流式回答考研问题，首字延迟 ≤ 2 秒'],
            ['辅助目标1', '语音提醒', '虚拟宠物每 25 分钟语音提醒休息，分心时即时提示'],
            ['辅助目标2', '学习数据记录', '记录专注时长、分心次数，生成每日学习报告'],
        ],
        col_widths=[2.5, 4, 9]
    )
    
    add_heading_styled(doc, '1.3 智能体定位', level=2)
    add_normal_para(doc, '本系统是一个多模态学习监督智能体，具备以下三方面能力：')
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('感知（Perception）：通过摄像头捕获用户面部图像，经 OpenCV/YOLOv8 检测人脸位置与状态')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('规划（Planning）：根据视觉感知结果和学习计时器状态，规划提醒策略（是否提醒、提醒内容、提醒方式）')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('行动（Action）：通过虚拟宠物动画和 Web Speech 语音合成执行提醒；通过大模型和 RAG 执行知识问答')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    add_heading_styled(doc, '二、技术选型', level=1)
    
    add_heading_styled(doc, '2.1 技术栈总览', level=2)
    add_table_simple(doc,
        ['层级', '技术方案', '版本', '选型理由'],
        [
            ['前端', 'HTML5 + Canvas + JavaScript', '-', '单文件部署，浏览器原生摄像头 API，Canvas 绘制检测框'],
            ['面部检测', 'OpenCV Haar Cascade + YOLOv8', '4.10.0 / 8.3.52', '后端检测，OpenCV 级联分类器为默认方案，YOLOv8 为增强方案'],
            ['语音合成', 'Web Speech API', '浏览器原生', '无需额外依赖，支持中文语音'],
            ['后端框架', 'Python FastAPI', '0.115.0', '异步高性能，自动生成 API 文档，AI/ML 生态友好'],
            ['大语言模型', 'DeepSeek API', '-', '兼容 OpenAI 格式，中文能力强，性价比高'],
            ['RAG 检索', '向量相似度 + 关键词回退', '-', 'pgvector 向量检索为主，SQLite LIKE 文本搜索为备选'],
            ['数据库', 'SQLite / PostgreSQL + pgvector', '-', '开发零配置，生产支持向量检索'],
            ['用户认证', 'JWT (python-jose)', '3.3.0', '无状态鉴权，前后端分离友好'],
            ['流式响应', 'SSE (sse-starlette)', '2.1.3', '大模型逐字流式输出，用户体验好'],
            ['HTTP 客户端', 'httpx', '≥0.27.0', '直接调用 DeepSeek API，避免 OpenAI SDK 代理兼容问题'],
            ['部署', 'Docker + Docker Compose', '-', '前端 Nginx + 后端 FastAPI + 数据库，一键部署'],
        ],
        col_widths=[2, 4, 2.5, 7]
    )
    
    add_heading_styled(doc, '2.2 核心依赖清单', level=2)
    add_code_block(doc, '''# 后端核心依赖
fastapi==0.115.0
uvicorn[standard]==0.30.6
sqlalchemy[asyncio]==2.0.35
python-jose[cryptography]==3.3.0
passlib[bcrypt]==1.7.4
pydantic==2.9.2
pydantic-settings==2.5.2
sse-starlette==2.1.3
aiosqlite==0.20.0
opencv-python-headless==4.10.0.84
ultralytics==8.3.52
httpx>=0.27.0''')
    
    add_heading_styled(doc, '三、系统架构设计', level=1)
    
    add_heading_styled(doc, '3.1 整体架构', level=2)
    add_normal_para(doc, '系统采用前后端分离架构，结合智能体感知-规划-行动的设计模式，整体分为四层：浏览器前端层、FastAPI 后端层、数据持久化层和外部 AI 服务层。')
    add_normal_para(doc, '前端负责摄像头采集、虚拟宠物展示和问答交互；后端负责视觉检测、问答服务和学习记录；数据层存储用户、对话、知识库和学习记录；外部服务提供大模型和向量能力。')
    
    add_heading_styled(doc, '3.2 模块职责', level=2)
    add_table_simple(doc,
        ['模块', '职责', '核心文件'],
        [
            ['用户认证模块', '注册、登录、JWT 签发与验证', 'app/routers/auth.py, app/core/security.py'],
            ['视觉检测模块', '图像解码、人脸检测、学习状态分类', 'app/routers/vision.py'],
            ['问答服务模块', 'DeepSeek API 调用、RAG 检索、SSE 流式输出', 'app/services/chat_service.py, app/services/rag_service.py'],
            ['对话管理模块', '会话创建、历史记录、消息存储', 'app/routers/chat.py'],
            ['学习记录模块', '学习计时、分心记录、报告生成', 'app/routers/study.py'],
            ['运行时设置', 'API Key 动态配置', 'app/routers/settings.py, app/core/runtime_config.py'],
            ['数据持久化', 'ORM 映射、数据库连接管理', 'app/core/database.py, app/models/'],
        ],
        col_widths=[3, 5.5, 7]
    )
    
    add_heading_styled(doc, '四、Agent 智能体设计', level=1)
    
    add_heading_styled(doc, '4.1 核心理念', level=2)
    add_normal_para(doc, '本系统的智能体遵循感知-规划-行动（Perception-Planning-Action）循环模式。感知层负责采集摄像头图像、检测人脸状态；规划层根据感知结果和规则引擎判定学习状态，决定是否提醒；行动层执行语音提醒、动画切换和问答响应。')
    
    add_heading_styled(doc, '4.2 感知层（Perception）', level=2)
    add_table_simple(doc,
        ['感知通道', '数据来源', '处理方式', '输出'],
        [
            ['视觉感知', '浏览器摄像头帧 (Base64 JPEG)', 'OpenCV Haar Cascade 多分类器融合检测', '人脸框坐标、置信度'],
            ['状态感知', '人脸位置与面积比', '规则引擎计算 (中心偏移、面积占比、宽高比)', '专注/分心/开小差/离开'],
            ['时间感知', '前端计时器', '25 分钟周期监测', '是否到达提醒时间'],
            ['文本感知', '用户输入的问题', 'DeepSeek API 向量化 + RAG 检索', '相关知识文档'],
        ],
        col_widths=[2.5, 3.5, 5, 4.5]
    )
    
    add_normal_para(doc, '视觉感知算法流程：', bold=True)
    steps = [
        '前端每 500ms 截取摄像头帧，压缩为 416px 宽 JPEG，Base64 编码',
        '后端解码图像，转灰度图，CLAHE 对比度增强',
        '加载多个 Haar Cascade 分类器（frontalface_default/alt2/alt/alt_tree）',
        '多尺度检测（scaleFactor: 1.03~1.05, minSize: 20~30px）',
        'IoU 合并重叠框（阈值 0.3），保留最大人脸',
        '计算人脸中心偏移量、面积占比，输出状态判定',
    ]
    for i, s in enumerate(steps):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(s)
        set_font(run, size=10.5)
        p.paragraph_format.line_spacing = 1.5
    
    add_heading_styled(doc, '4.3 规划层（Planning）', level=2)
    add_normal_para(doc, '状态判定规则（优先级从高到低）：', bold=True)
    add_table_simple(doc,
        ['条件', '状态', '专注分数', '原因'],
        [
            ['无人脸或面积比 < 1.5%', '离开', '10', '人脸太小或离屏幕太远'],
            ['中心偏移 X > 62% 或 Y > 72%', '开小差', '≤ 32', '人脸偏离画面中心'],
            ['检测置信度 < 0.4', '分心', '≤ 55', '检测置信度偏低'],
            ['中心偏移 X < 50% 且 Y < 62% 且面积 ≥ 2%', '专注', '≥ 78', '检测到稳定人脸'],
            ['其他情况', '分心', '≤ 65', '人脸位置不稳定'],
        ],
        col_widths=[5.5, 2, 2, 5.5]
    )
    
    add_normal_para(doc, '提醒策略规划：', bold=True)
    add_table_simple(doc,
        ['触发条件', '提醒类型', '提醒内容'],
        [
            ['25 分钟计时到达', '休息提醒', '"你已经学习25分钟啦，起来走动一下吧~"'],
            ['检测到"分心"状态', '分心提醒', '"注意力好像飘走了哦，深呼吸重新集中一下吧~"'],
            ['检测到"开小差"状态', '开小差提醒', '"好像走神了呢，要继续学习吗？"'],
            ['检测到"离开"状态超 30 秒', '离开提醒', '"你去哪里了呀？学习还没结束哦~"'],
        ],
        col_widths=[4.5, 3, 7.5]
    )
    
    add_heading_styled(doc, '4.4 行动层（Action）', level=2)
    add_table_simple(doc,
        ['行动类型', '执行方式', '技术实现'],
        [
            ['语音提醒', 'Web Speech API 合成中文语音', 'SpeechSynthesisUtterance, lang=zh-CN, rate=0.9'],
            ['宠物动画', 'CSS 动画切换宠物表情与动作', '专注→安静看书；分心→敲屏幕；离开→张望'],
            ['人脸框绘制', 'Canvas 实时绘制检测框与状态标签', '绿色框(专注) / 橙色框(开小差) / 红色遮罩(分心)'],
            ['知识问答', 'DeepSeek API 流式生成回答', 'SSE 推送，逐字显示，末尾附知识总结'],
        ],
        col_widths=[3, 4.5, 7.5]
    )
    
    add_heading_styled(doc, '4.5 问答工作流', level=2)
    qa_steps = [
        ('问题向量化', '调用 DeepSeek embedding API 将用户问题转为向量；失败则返回模拟向量'),
        ('RAG 知识检索', '生产环境使用 pgvector 余弦相似度搜索 top-5 相关文档；开发环境使用 SQLite LIKE 关键词搜索作为回退'),
        ('构建 Prompt', '组装科目专属系统提示，注入检索到的参考资料，追加知识总结要求'),
        ('DeepSeek API 调用', '使用 httpx 原生流式调用，逐字 yield 给 SSE，max_tokens=4096, temperature=0.7'),
        ('保存对话历史', '用户消息和 AI 回复存入 chat_history 表，通过 session_id 关联会话'),
    ]
    for i, (title, desc) in enumerate(qa_steps):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'{title}：')
        set_font(run, size=10.5, bold=True)
        run2 = p.add_run(desc)
        set_font(run2, size=10.5)
        p.paragraph_format.line_spacing = 1.5
    
    add_heading_styled(doc, '五、数据模型设计', level=1)
    
    add_heading_styled(doc, '5.1 数据库表结构', level=2)
    
    add_heading_styled(doc, '用户表 (users)', level=3)
    add_table_simple(doc,
        ['字段', '类型', '约束', '说明'],
        [
            ['id', 'Integer', 'PK, 自增', '用户ID'],
            ['username', 'String(50)', 'UNIQUE, NOT NULL', '用户名'],
            ['email', 'String(100)', 'UNIQUE, NOT NULL', '邮箱'],
            ['hashed_password', 'String(255)', 'NOT NULL', 'bcrypt 哈希密码'],
            ['created_at', 'DateTime', 'DEFAULT NOW()', '创建时间'],
        ],
        col_widths=[3.5, 3, 3.5, 5]
    )
    
    add_heading_styled(doc, '对话历史表 (chat_history)', level=3)
    add_table_simple(doc,
        ['字段', '类型', '约束', '说明'],
        [
            ['id', 'Integer', 'PK, 自增', '消息ID'],
            ['user_id', 'Integer', 'FK(users.id)', '用户ID'],
            ['session_id', 'String(36)', 'NOT NULL, INDEX', '会话ID (UUID)'],
            ['subject', 'String(20)', 'NOT NULL', '科目'],
            ['role', 'String(20)', 'NOT NULL', '角色 (user/assistant)'],
            ['content', 'Text', 'NOT NULL', '消息内容'],
            ['created_at', 'DateTime', 'DEFAULT NOW()', '创建时间'],
        ],
        col_widths=[3, 3, 3.5, 5.5]
    )
    
    add_heading_styled(doc, '知识库表 (knowledge_base)', level=3)
    add_table_simple(doc,
        ['字段', '类型', '约束', '说明'],
        [
            ['id', 'Integer', 'PK, 自增', '知识ID'],
            ['subject', 'String(50)', 'NOT NULL, INDEX', '科目'],
            ['content', 'Text', 'NOT NULL', '知识内容'],
            ['embedding', 'VECTOR(1536) / Text', 'NULLABLE', '向量 (PG) / 空 (SQLite)'],
            ['source', 'String(255)', 'NULLABLE', '来源'],
            ['created_at', 'DateTime', 'DEFAULT NOW()', '创建时间'],
        ],
        col_widths=[3, 3.5, 3, 5.5]
    )
    
    add_heading_styled(doc, '学习记录表 (study_sessions)', level=3)
    add_table_simple(doc,
        ['字段', '类型', '约束', '说明'],
        [
            ['id', 'Integer', 'PK, 自增', '记录ID'],
            ['user_id', 'Integer', 'FK(users.id)', '用户ID'],
            ['start_time', 'DateTime', 'DEFAULT NOW()', '开始时间'],
            ['end_time', 'DateTime', 'NULLABLE', '结束时间'],
            ['focus_duration_seconds', 'Integer', 'DEFAULT 0', '专注时长(秒)'],
            ['distraction_count', 'Integer', 'DEFAULT 0', '分心次数'],
        ],
        col_widths=[3.5, 3, 3, 5.5]
    )
    
    add_heading_styled(doc, '六、安全设计', level=1)
    
    add_heading_styled(doc, '6.1 认证与授权', level=2)
    add_table_simple(doc,
        ['安全措施', '实现方式'],
        [
            ['密码存储', 'bcrypt 哈希 (passlib)'],
            ['Token 签发', 'JWT HS256, 有效期 24 小时'],
            ['Token 验证', '每个受保护端点通过 get_current_user 依赖注入验证'],
            ['API Key 保护', '运行时内存存储，不写入文件；通过 settings API 动态配置'],
        ],
        col_widths=[4, 11]
    )
    
    add_heading_styled(doc, '6.2 数据安全', level=2)
    add_table_simple(doc,
        ['安全措施', '实现方式'],
        [
            ['SQL 注入防护', 'SQLAlchemy ORM 参数化查询'],
            ['CORS 配置', '默认允许所有源 (开发)，生产环境限制具体域名'],
            ['代理环境变量清除', '启动时清除 HTTP_PROXY/HTTPS_PROXY 等，避免 API 调用泄露'],
            ['摄像头隐私', '图像仅在内存中处理检测，不持久化存储用户面部图像'],
        ],
        col_widths=[4, 11]
    )
    
    add_heading_styled(doc, '七、部署方案', level=1)
    
    add_heading_styled(doc, '7.1 开发环境', level=2)
    add_code_block(doc, '''前端: Python http.server 8080 端口
后端: Uvicorn 8000 端口, SQLite 数据库
数据库: SQLite (ai_tutor.db, 零配置)''')
    
    add_heading_styled(doc, '7.2 生产环境 (Docker Compose)', level=2)
    add_code_block(doc, '''version: '3.8'
services:
  frontend:
    build: ./ai-tutor-frontend
    ports: ["80:80"]
    depends_on: [backend]

  backend:
    build: ./ai-tutor-backend
    ports: ["8000:8000"]
    environment:
      - DATABASE_URL=postgresql+asyncpg://user:pass@db:5432/ai_tutor
      - DEEPSEEK_API_KEY=${DEEPSEEK_API_KEY}
    depends_on: [db]

  db:
    image: pgvector/pgvector:pg16
    environment:
      - POSTGRES_USER=user
      - POSTGRES_PASSWORD=pass
      - POSTGRES_DB=ai_tutor
    volumes: [pgdata:/var/lib/postgresql/data]

volumes:
  pgdata:''')
    
    add_normal_para(doc, '注意：摄像头 API 只能在 localhost 或 HTTPS 下调用，生产环境需配置 HTTPS (Let\'s Encrypt + Nginx 反向代理)。')
    
    add_heading_styled(doc, '八、项目里程碑', level=1)
    add_table_simple(doc,
        ['阶段', '内容', '交付物'],
        [
            ['第1阶段', '环境搭建、数据库设计', '项目骨架、数据库表结构'],
            ['第2阶段', '前端 UI（摄像头+宠物+对话）', '可交互的前端界面'],
            ['第3阶段', '后端 API（认证+问答+学习记录）', '可调用的 REST API'],
            ['第4阶段', '大模型集成 + RAG 知识库', '智能问答功能上线'],
            ['第5阶段', '视觉检测集成（OpenCV+YOLOv8）', '学习状态监督功能上线'],
            ['第6阶段', '联调测试、Docker 部署', '可部署的完整系统'],
        ],
        col_widths=[2.5, 5, 7.5]
    )
    
    add_heading_styled(doc, '九、风险与应对', level=1)
    add_table_simple(doc,
        ['风险', '概率', '影响', '应对措施'],
        [
            ['DeepSeek API 不可用', '中', '高', '开发模式返回模拟回答；预留 OpenAI 备用配置'],
            ['OpenCV 检测精度不足', '中', '中', '多分类器融合 + CLAHE 增强；预留 YOLOv8 模型接口'],
            ['浏览器不支持摄像头', '低', '高', '降级为纯计时器模式，无视觉检测'],
            ['代理环境变量冲突', '中', '中', '三重清除机制（启动时/导入时/客户端创建时）'],
            ['SQLite 并发限制', '低', '低', '生产环境切换 PostgreSQL + pgvector'],
        ],
        col_widths=[4, 1.5, 1.5, 8]
    )
    
    output_path = r'D:\AI Study\.trae-cn\worktrees\AI Study\feat-yolov8-face-detection-KcMOL5\docs\01-智能体项目立项与架构说明书.docx'
    doc.save(output_path)
    print(f'文档1已生成: {output_path}')


def generate_doc2():
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import nsmap
    from docx.oxml import parse_xml
    
    doc = Document()
    
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI 学习助手 — 系统整体架构图')
    set_font(run, size=20, bold=True, color=(31, 73, 125))
    title.paragraph_format.space_after = Pt(6)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('考研助手智能体 | 前后端分离 + Agent 感知-规划-行动架构')
    set_font(run, size=11, color=(100, 116, 139))
    subtitle.paragraph_format.space_after = Pt(18)
    
    doc.add_paragraph()
    
    def add_layer(title_text, color_hex, bg_hex, modules, title_color=(31, 73, 125)):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        cell.width = Cm(24)
        set_cell_shading(cell, bg_hex)
        
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        set_font(run, size=12, bold=True, color=title_color)
        p.paragraph_format.space_after = Pt(8)
        
        inner_table = cell.add_table(rows=1, cols=len(modules))
        inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (m_title, m_desc) in enumerate(modules):
            m_cell = inner_table.rows[0].cells[i]
            m_cell.width = Cm(7.5)
            set_cell_shading(m_cell, 'FFFFFF')
            
            m_cell.text = ''
            mp = m_cell.paragraphs[0]
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mr = mp.add_run(m_title)
            set_font(mr, size=11, bold=True, color=title_color)
            mp.paragraph_format.space_after = Pt(4)
            
            for line in m_desc:
                lp = m_cell.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = lp.add_run(line)
                set_font(lr, size=9, color=(71, 85, 105))
                lp.paragraph_format.line_spacing = 1.3
        
        doc.add_paragraph()
        return table
    
    frontend_modules = [
        ('摄像头模块', ['getUserMedia 采集', 'Canvas 绘制人脸框', '500ms 截帧发送']),
        ('虚拟宠物 UI', ['CSS 动画表情切换', 'Web Speech 语音提醒', '25分钟定时器']),
        ('问答对话界面', ['Markdown 渲染', '会话列表管理', '科目切换隔离']),
    ]
    add_layer('浏览器前端层 (Port 8080)', '3B82F6', 'EFF6FF', frontend_modules, title_color=(29, 78, 216))
    
    arrow_p = doc.add_paragraph()
    arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = arrow_p.add_run('▲▼  HTTP / fetch + SSE (Server-Sent Events)  ▲▼')
    set_font(ar, size=10, color=(100, 116, 139))
    doc.add_paragraph()
    
    backend_modules = [
        ('视觉检测模块 (Agent 感知)', ['OpenCV Haar Cascade', '多分类器融合检测', 'YOLOv8 人脸检测(增强)', '状态分类: 专注/分心/开小差/离开']),
        ('问答服务模块 (Agent 行动)', ['DeepSeek API 流式调用', 'RAG 知识库检索', '科目专属 Prompt', 'SSE 逐字推送']),
        ('学习记录模块 (数据持久化)', ['学习会话管理', '专注时长统计', '分心次数记录', '每日学习报告']),
    ]
    add_layer('FastAPI 后端层 (Port 8000)', '10B981', 'ECFDF5', backend_modules, title_color=(4, 120, 87))
    
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = auth_p.add_run('🔐 JWT 认证中间件 — 注册 / 登录 / Token 验证 / get_current_user 依赖注入')
    set_font(ar, size=10, bold=True, color=(146, 64, 14))
    auth_p.paragraph_format.space_before = Pt(4)
    auth_p.paragraph_format.space_after = Pt(4)
    
    arrow_p = doc.add_paragraph()
    arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = arrow_p.add_run('▲▼  SQLAlchemy ORM + AsyncSession  ▲▼')
    set_font(ar, size=10, color=(100, 116, 139))
    doc.add_paragraph()
    
    data_modules = [
        ('users 表', ['id, username, email', 'hashed_password', 'created_at']),
        ('chat_history 表', ['id, user_id, session_id', 'subject, role, content', 'created_at']),
        ('knowledge_base 表', ['id, subject, content', 'embedding(向量)', 'source']),
        ('study_sessions 表', ['id, user_id, start_time', 'end_time, focus_duration', 'distraction_count']),
    ]
    add_layer('数据持久化层', 'F59E0B', 'FFFBEB', data_modules, title_color=(180, 83, 9))
    
    env_p = doc.add_paragraph()
    env_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = env_p.add_run('开发环境: SQLite (ai_tutor.db)   |   生产环境: PostgreSQL + pgvector 向量检索')
    set_font(er, size=9, color=(146, 64, 14))
    doc.add_paragraph()
    
    arrow_p = doc.add_paragraph()
    arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = arrow_p.add_run('▲▼  httpx 原生 HTTP 调用 (兼容 OpenAI 格式)  ▲▼')
    set_font(ar, size=10, color=(100, 116, 139))
    doc.add_paragraph()
    
    external_modules = [
        ('DeepSeek API (deepseek-chat)', ['模型: deepseek-chat', '用途: 考研问答主模型', '流式输出 (stream=True)', 'max_tokens=4096']),
        ('DeepSeek Embedding', ['模型: deepseek-embedding', '用途: RAG 文本向量化', '维度: 1536', '失败回退: 关键词搜索']),
        ('', ['', '', '', '']),
    ]
    add_layer('外部 AI 服务', '8B5CF6', 'F5F3FF', external_modules, title_color=(109, 40, 217))
    
    doc.add_paragraph()
    legend_title = doc.add_paragraph()
    legend_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = legend_title.add_run('图例说明')
    set_font(lr, size=11, bold=True, color=(71, 85, 105))
    
    legend_table = doc.add_table(rows=1, cols=5)
    legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    legend_items = [
        ('3B82F6', '前端层'),
        ('10B981', '后端层'),
        ('F59E0B', '数据层'),
        ('8B5CF6', '外部服务'),
        ('FEF3C7', '安全/中间件'),
    ]
    for i, (color, name) in enumerate(legend_items):
        cell = legend_table.rows[0].cells[i]
        cell.width = Cm(4.5)
        set_cell_shading(cell, color)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        set_font(r, size=10, bold=True, color=(71, 85, 105) if color == 'FEF3C7' else (255, 255, 255))
    
    output_path = r'D:\AI Study\.trae-cn\worktrees\AI Study\feat-yolov8-face-detection-KcMOL5\docs\02-系统整体架构图.docx'
    doc.save(output_path)
    print(f'文档2已生成: {output_path}')


def generate_doc3():
    from docx.enum.section import WD_ORIENT
    
    doc = Document()
    
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI 学习助手 — Agent 工作流设计图')
    set_font(run, size=20, bold=True, color=(31, 73, 125))
    title.paragraph_format.space_after = Pt(6)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('感知-规划-行动循环 | 视觉监督 Agent + 问答 Agent 双工作流')
    set_font(run, size=11, color=(100, 116, 139))
    subtitle.paragraph_format.space_after = Pt(18)
    
    loop_title = doc.add_paragraph()
    loop_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = loop_title.add_run('Agent 核心循环 (Perception-Planning-Action)')
    set_font(lr, size=14, bold=True, color=(15, 23, 42))
    loop_title.paragraph_format.space_after = Pt(16)
    
    loop_table = doc.add_table(rows=1, cols=5)
    loop_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_widths = [5.5, 1.5, 5.5, 1.5, 5.5]
    
    stages = [
        ('感知', 'Perception', 'EFF6FF', (29, 78, 216), ['📷 摄像头帧采集 (500ms)', '🔍 OpenCV 人脸检测', '🧠 YOLOv8 状态分类', '⏱️ 学习计时器状态', '📝 用户文本输入']),
        ('规划', 'Planning', 'FFFBEB', (180, 83, 9), ['📐 人脸中心偏移计算', '📊 面积占比 & 宽高比', '🎯 状态判定规则引擎', '⏰ 25分钟提醒决策', '📚 RAG 知识检索策略']),
        ('行动', 'Action', 'ECFDF5', (4, 120, 87), ['🔊 Web Speech 语音提醒', '🦉 宠物动画切换', '🖼️ Canvas 人脸框绘制', '🤖 DeepSeek 流式问答', '💾 对话历史持久化']),
    ]
    
    for i in range(3):
        idx = i * 2
        cell = loop_table.rows[0].cells[idx]
        cell.width = Cm(cell_widths[idx])
        set_cell_shading(cell, stages[i][2])
        
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(stages[i][0])
        set_font(r, size=14, bold=True, color=stages[i][3])
        p.paragraph_format.space_after = Pt(2)
        
        ep = cell.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        er = ep.add_run(stages[i][1])
        set_font(er, size=9, color=(100, 116, 139))
        ep.paragraph_format.space_after = Pt(8)
        
        for item in stages[i][4]:
            ip = cell.add_paragraph()
            ip.paragraph_format.left_indent = Cm(0.3)
            ir = ip.add_run(item)
            set_font(ir, size=9.5)
            ip.paragraph_format.line_spacing = 1.4
        
        if i < 2:
            arrow_cell = loop_table.rows[0].cells[idx + 1]
            arrow_cell.width = Cm(cell_widths[idx + 1])
            arrow_cell.text = ''
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arun = ap.add_run('→')
            set_font(arun, size=24, color=(148, 163, 184))
    
    loop_back = doc.add_paragraph()
    loop_back.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lb = loop_back.add_run('↻ 循环回到感知层 (每 500ms 下一周期)')
    set_font(lb, size=10, color=(100, 116, 139))
    loop_back.paragraph_format.space_before = Pt(8)
    loop_back.paragraph_format.space_after = Pt(20)
    
    wf_title = doc.add_paragraph()
    wf_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wr = wf_title.add_run('两条工作流详情')
    set_font(wr, size=14, bold=True, color=(15, 23, 42))
    wf_title.paragraph_format.space_after = Pt(12)
    
    wf_table = doc.add_table(rows=1, cols=2)
    wf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    vision_steps = [
        ('摄像头帧采集', '前端 getUserMedia 获取视频流，Canvas 截帧压缩为 416px JPEG，Base64 编码'),
        ('后端图像解码 + 预处理', 'cv2.imdecode 解码 → BGR 转灰度 → equalizeHist 直方图均衡化 → CLAHE 自适应对比度增强'),
        ('多分类器人脸检测', '加载 4 个 Haar Cascade 分类器 (default/alt2/alt/alt_tree)，多尺度检测 (scaleFactor 1.03~1.05)，IoU 合并重叠框'),
        ('状态判定 (规则引擎)', '计算人脸中心偏移量 (center_dx/dy)、面积占比 (area_ratio)、置信度 (score)，按优先级判定学习状态'),
        ('前端行动执行', 'Canvas 绘制人脸框 (绿/橙/红) → 宠物表情切换 → 语音提醒触发 → 分心计数 +1'),
    ]
    
    chat_steps = [
        ('用户提问 + 历史加载', '前端发送 question + subject + session_id → 后端从 chat_history 加载该会话历史消息'),
        ('问题向量化', '调用 DeepSeek embedding API 将问题转为 1536 维向量；失败则返回模拟向量'),
        ('RAG 知识检索', 'pgvector 余弦相似度搜索 top-5 相关文档 (生产) / SQLite LIKE 关键词搜索 (开发回退)'),
        ('Prompt 构建', '科目专属系统提示 (政治/英语/数学/专业课) + 注入检索参考资料 + 追加知识总结要求'),
        ('DeepSeek 流式生成', 'httpx 原生 stream 调用，逐 chunk yield 给 SSE，前端逐字渲染'),
        ('对话历史持久化', '用户消息 + AI 完整回复存入 chat_history 表，session_id 关联会话'),
    ]
    
    def build_workflow_cell(cell, title, color_hex, title_color, steps, step_color):
        cell.width = Cm(12)
        set_cell_shading(cell, 'F8FAFC')
        
        cell.text = ''
        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(title)
        set_font(hr, size=13, bold=True, color=title_color)
        hp.paragraph_format.space_after = Pt(10)
        
        for i, (s_title, s_desc) in enumerate(steps):
            step_table = cell.add_table(rows=1, cols=2)
            step_table.autofit = False
            
            num_cell = step_table.rows[0].cells[0]
            num_cell.width = Cm(1)
            set_cell_shading(num_cell, step_color)
            num_cell.text = ''
            np = num_cell.paragraphs[0]
            np.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nr = np.add_run(str(i + 1))
            set_font(nr, size=11, bold=True, color=(255, 255, 255))
            
            content_cell = step_table.rows[0].cells[1]
            content_cell.width = Cm(10.5)
            content_cell.text = ''
            cp = content_cell.paragraphs[0]
            cr = cp.add_run(s_title)
            set_font(cr, size=10.5, bold=True)
            cp.paragraph_format.space_after = Pt(2)
            
            dp = content_cell.add_paragraph()
            dr = dp.add_run(s_desc)
            set_font(dr, size=9, color=(71, 85, 105))
            dp.paragraph_format.line_spacing = 1.4
            
            if i < len(steps) - 1:
                ap = cell.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ar = ap.add_run('│')
                set_font(ar, size=12, color=(203, 213, 225))
                ap.paragraph_format.space_before = Pt(0)
                ap.paragraph_format.space_after = Pt(0)
    
    build_workflow_cell(wf_table.rows[0].cells[0], '工作流 A：视觉学习监督', 
                       'EFF6FF', (29, 78, 216), vision_steps, '3B82F6')
    build_workflow_cell(wf_table.rows[0].cells[1], '工作流 B：考研智能问答',
                       'ECFDF5', (4, 120, 87), chat_steps, '10B981')
    
    doc.add_paragraph()
    decision_title = doc.add_paragraph()
    decision_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = decision_title.add_run('学习状态判定决策表 (规划层规则引擎)')
    set_font(dr, size=14, bold=True, color=(15, 23, 42))
    decision_title.paragraph_format.space_before = Pt(16)
    decision_title.paragraph_format.space_after = Pt(10)
    
    add_table_simple(doc,
        ['优先级', '触发条件', '判定状态', '专注分数', '行动响应'],
        [
            ['1 (最高)', '无人脸 / 面积比 < 1.5%', '离开', '10', '红色遮罩 + "你去哪里了呀？"语音'],
            ['2', '中心偏移 X > 62% 或 Y > 72%', '开小差', '≤ 32', '橙色框 + "好像走神了呢"语音'],
            ['3', '检测置信度 < 0.4', '分心', '≤ 55', '红色框 + "注意力飘走了哦"语音'],
            ['4', '偏移 X<50% 且 Y<62% 且 面积≥2%', '专注', '≥ 78', '绿色框 + 宠物安静看书动画'],
            ['5 (默认)', '其他情况', '分心', '≤ 65', '红色框 + 分心提醒'],
            ['定时器', '25分钟计时到达', '—', '—', '宠物跳出 + "该休息一下啦~"语音'],
        ],
        col_widths=[2, 6, 2.5, 2, 7]
    )
    
    output_path = r'D:\AI Study\.trae-cn\worktrees\AI Study\feat-yolov8-face-detection-KcMOL5\docs\03-Agent工作流设计图.docx'
    doc.save(output_path)
    print(f'文档3已生成: {output_path}')


def generate_doc4():
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('前后端 API 接口定义表')
    set_font(run, size=22, bold=True, color=(31, 73, 125))
    title.paragraph_format.space_after = Pt(12)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'
    info_data = [
        ('项目名称', 'AI 学习助手（考研助手）'),
        ('文档版本', 'V1.0'),
        ('编写日期', '2026-07-03'),
        ('Base URL', 'http://localhost:8000'),
        ('认证方式', 'Bearer Token (JWT)'),
    ]
    for i, (k, v) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = ''
        p1 = row[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_font(r1, size=10.5, bold=True)
        set_cell_shading(row[0], 'D9E2F3')
        
        row[1].text = ''
        p2 = row[1].paragraphs[0]
        r2 = p2.add_run(v)
        set_font(r2, size=10.5)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '一、接口总览', level=1)
    
    overview_data = [
        ['1', 'GET', '/api/health', '否', '系统', '健康检查'],
        ['2', 'POST', '/api/auth/register', '否', '认证', '用户注册'],
        ['3', 'POST', '/api/auth/login', '否', '认证', '用户登录'],
        ['4', 'GET', '/api/auth/me', '是', '认证', '获取当前用户'],
        ['5', 'POST', '/api/chat/sessions', '是', '问答', '创建会话'],
        ['6', 'POST', '/api/chat/send', '是', '问答', '发送消息（非流式）'],
        ['7', 'POST', '/api/chat/stream', '是', '问答', 'SSE 流式对话'],
        ['8', 'GET', '/api/chat/history/{session_id}', '是', '问答', '获取对话历史'],
        ['9', 'GET', '/api/chat/sessions', '是', '问答', '获取会话列表'],
        ['10', 'DELETE', '/api/chat/history/{session_id}', '是', '问答', '删除对话'],
        ['11', 'GET', '/api/chat/debug/info', '否', '问答', '调试信息'],
        ['12', 'POST', '/api/vision/detect', '否', '视觉', '人脸检测与状态识别'],
        ['13', 'POST', '/api/study/session/start', '是', '学习', '开始学习会话'],
        ['14', 'POST', '/api/study/session/end', '是', '学习', '结束学习会话'],
        ['15', 'POST', '/api/study/distraction', '是', '学习', '记录分心事件'],
        ['16', 'GET', '/api/study/report', '是', '学习', '获取每日学习报告'],
        ['17', 'POST', '/api/settings/apikey', '是', '设置', '设置 API Key'],
        ['18', 'GET', '/api/settings/apikey/{key_name}', '是', '设置', '查询 API Key 状态'],
    ]
    add_table_simple(doc,
        ['序号', '方法', '路径', '认证', '模块', '说明'],
        overview_data,
        col_widths=[1.2, 1.8, 5, 1.2, 1.5, 4.3]
    )
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '二、通用说明', level=1)
    
    add_heading_styled(doc, '2.1 请求格式', level=2)
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Content-Type: application/json（除视觉检测的图像数据外）')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('所有受保护的接口需在 Header 中携带：Authorization: Bearer {access_token}')
    set_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.5
    
    add_heading_styled(doc, '2.2 HTTP 状态码', level=2)
    add_table_simple(doc,
        ['状态码', '含义'],
        [
            ['200', '请求成功'],
            ['201', '创建成功'],
            ['400', '请求参数错误'],
            ['401', '未认证 / Token 无效'],
            ['404', '资源不存在'],
            ['500', '服务器内部错误'],
        ],
        col_widths=[3, 12]
    )
    
    add_heading_styled(doc, '2.3 错误响应格式', level=2)
    add_code_block(doc, '''{
  "error": "错误类型",
  "message": "错误详情",
  "path": "请求路径"
}''')
    
    add_heading_styled(doc, '三、用户认证模块', level=1)
    
    add_heading_styled(doc, '3.1 用户注册', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/auth/register'],
            ['认证', '不需要'],
            ['说明', '注册新用户，返回 JWT Token'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '约束', '说明'],
        [
            ['username', 'string', '是', '2-50 字符', '用户名'],
            ['email', 'string', '是', '合法邮箱格式', '邮箱'],
            ['password', 'string', '是', '≥ 6 位', '密码'],
        ],
        col_widths=[2.5, 2, 1.5, 3.5, 5.5]
    )
    
    add_normal_para(doc, '请求示例：', bold=True, indent=False)
    add_code_block(doc, '''{
  "username": "student01",
  "email": "student@example.com",
  "password": "123456"
}''')
    
    add_normal_para(doc, '成功响应 (201)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "access_token": "eyJhbGciOiJIUzI1NiIs...",
  "token_type": "bearer",
  "user": {
    "id": 1,
    "username": "student01",
    "email": "student@example.com",
    "created_at": "2026-07-03T10:00:00"
  }
}''')
    
    add_normal_para(doc, '错误响应：', bold=True, indent=False)
    add_table_simple(doc,
        ['状态码', 'detail'],
        [
            ['400', '用户名已被注册'],
            ['400', '邮箱已被注册'],
        ],
        col_widths=[3, 12]
    )
    
    add_heading_styled(doc, '3.2 用户登录', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/auth/login'],
            ['认证', '不需要'],
            ['说明', '用户登录，返回 JWT Token'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['username', 'string', '是', '用户名'],
            ['password', 'string', '是', '密码'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "access_token": "eyJhbGciOiJIUzI1NiIs...",
  "token_type": "bearer",
  "user": {
    "id": 1,
    "username": "student01",
    "email": "student@example.com",
    "created_at": "2026-07-03T10:00:00"
  }
}''')
    
    add_heading_styled(doc, '3.3 获取当前用户信息', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'GET'],
            ['路径', '/api/auth/me'],
            ['认证', '需要'],
            ['说明', '获取当前登录用户的详细信息'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "id": 1,
  "username": "student01",
  "email": "student@example.com",
  "created_at": "2026-07-03T10:00:00"
}''')
    
    add_heading_styled(doc, '四、考研问答模块', level=1)
    
    add_heading_styled(doc, '4.1 创建对话会话', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/chat/sessions'],
            ['认证', '需要'],
            ['说明', '创建一个新的对话会话'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['subject', 'string', '是', '科目（政治/英语/数学/专业课）'],
            ['title', 'string | null', '否', '会话标题，默认"新对话"'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (201)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "session_id": "550e8400-e29b-41d4-a716-446655440000",
  "subject": "数学",
  "title": "高数极限问题",
  "messages": []
}''')
    
    add_heading_styled(doc, '4.2 发送消息（非流式）', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/chat/send'],
            ['认证', '需要'],
            ['说明', '发送问题，等待完整回答后返回'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['question', 'string', '是', '用户问题'],
            ['subject', 'string', '是', '科目（默认"数学"）'],
            ['session_id', 'string | null', '否', '会话ID，为空则自动生成'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "session_id": "550e8400-e29b-41d4-a716-446655440000",
  "reply": "洛必达法则是用于求解不定式极限的重要方法..."
}''')
    
    add_heading_styled(doc, '4.3 流式对话（SSE）', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/chat/stream'],
            ['认证', '需要'],
            ['说明', 'SSE 流式对话，逐字返回 AI 回答'],
            ['响应类型', 'text/event-stream'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['question', 'string', '是', '用户问题'],
            ['subject', 'string', '是', '科目'],
            ['session_id', 'string | null', '否', '会话ID'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, 'SSE 事件格式：', bold=True, indent=False)
    add_code_block(doc, '''event: message
data: {"type": "text", "content": "洛必达"}

event: message
data: {"type": "text", "content": "法则是"}

event: message
data: {"type": "done", "session_id": "550e8400-..."}''')
    
    add_normal_para(doc, 'SSE 事件类型：', bold=True, indent=False)
    add_table_simple(doc,
        ['type', '说明', 'content 字段'],
        [
            ['text', '文本片段', 'AI 回答的一个 chunk'],
            ['done', '流结束', 'session_id（完整会话ID）'],
            ['error', '错误', '错误信息'],
        ],
        col_widths=[2.5, 4, 8.5]
    )
    
    add_heading_styled(doc, '4.4 获取对话历史', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'GET'],
            ['路径', '/api/chat/history/{session_id}'],
            ['认证', '需要'],
            ['说明', '获取指定会话的完整对话记录'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '路径参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['参数', '类型', '说明'],
        [
            ['session_id', 'string', '会话ID (UUID)'],
        ],
        col_widths=[3, 3, 9]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "session_id": "550e8400-e29b-41d4-a716-446655440000",
  "subject": "数学",
  "messages": [
    {
      "id": 1,
      "role": "user",
      "content": "什么是洛必达法则？"
    },
    {
      "id": 2,
      "role": "assistant",
      "content": "洛必达法则是用于求解不定式极限的重要方法..."
    }
  ]
}''')
    
    add_heading_styled(doc, '4.5 获取所有会话列表', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'GET'],
            ['路径', '/api/chat/sessions'],
            ['认证', '需要'],
            ['说明', '获取当前用户的所有对话会话，按最后活跃时间排序'],
        ],
        col_widths=[3, 12]
    )
    
    add_heading_styled(doc, '4.6 删除对话', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'DELETE'],
            ['路径', '/api/chat/history/{session_id}'],
            ['认证', '需要'],
            ['说明', '删除指定会话的所有对话记录'],
        ],
        col_widths=[3, 12]
    )
    
    add_heading_styled(doc, '五、视觉检测模块', level=1)
    
    add_heading_styled(doc, '5.1 人脸检测与状态识别', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/vision/detect'],
            ['认证', '不需要'],
            ['说明', '接收摄像头帧图像，返回人脸检测结果和学习状态'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['image', 'string', '是', 'Base64 编码的 JPEG 图像，支持 data URL 前缀'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "status": "专注",
  "focus_score": 82,
  "reason": "检测到稳定人脸，视线集中",
  "faces": [
    {
      "bbox": [120.5, 60.3, 200.8, 280.4],
      "score": 0.80,
      "source": "opencv"
    }
  ],
  "model": "opencv-fallback",
  "frame": {
    "width": 416,
    "height": 312
  }
}''')
    
    add_normal_para(doc, '字段说明：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '说明'],
        [
            ['status', 'string', '学习状态：专注 / 分心 / 开小差 / 离开'],
            ['focus_score', 'int', '专注分数 (0-100)'],
            ['reason', 'string', '状态判定原因'],
            ['faces', 'array', '人脸检测结果列表'],
            ['faces[].bbox', 'array[float]', '人脸框 [x, y, width, height]'],
            ['faces[].score', 'float', '检测置信度 (0-1)'],
            ['faces[].source', 'string', '检测来源：opencv / yolov8'],
            ['model', 'string', '使用的模型：yolov8 / opencv-fallback / unavailable'],
            ['frame', 'object', '图像尺寸信息'],
        ],
        col_widths=[3.5, 2.5, 9]
    )
    
    add_normal_para(doc, '状态说明：', bold=True, indent=False)
    add_table_simple(doc,
        ['status', '含义', 'focus_score 范围', '触发条件'],
        [
            ['专注', '用户正在专注学习', '78-100', '人脸居中且面积 ≥ 2%'],
            ['分心', '用户注意力不集中', '0-65', '检测置信度低或位置不稳定'],
            ['开小差', '用户明显走神', '0-32', '人脸偏离画面中心'],
            ['离开', '用户不在画面中', '0-10', '无人脸或面积 < 1.5%'],
        ],
        col_widths=[2.5, 4, 3.5, 5]
    )
    
    add_heading_styled(doc, '六、学习记录模块', level=1)
    
    add_heading_styled(doc, '6.1 开始学习会话', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/study/session/start'],
            ['认证', '需要'],
            ['说明', '开始一次学习记录会话'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '成功响应 (201)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "id": 1,
  "start_time": "2026-07-03T10:00:00",
  "end_time": null,
  "focus_duration_seconds": 0,
  "distraction_count": 0
}''')
    
    add_heading_styled(doc, '6.2 结束学习会话', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/study/session/end'],
            ['认证', '需要'],
            ['说明', '结束学习会话，自动计算专注时长'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['session_id', 'int', '是', '学习会话ID'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "id": 1,
  "start_time": "2026-07-03T10:00:00",
  "end_time": "2026-07-03T10:50:00",
  "focus_duration_seconds": 3000,
  "distraction_count": 3
}''')
    
    add_heading_styled(doc, '6.3 记录分心事件', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/study/distraction'],
            ['认证', '需要'],
            ['说明', '记录一次分心事件，分心计数 +1'],
        ],
        col_widths=[3, 12]
    )
    
    add_heading_styled(doc, '6.4 获取每日学习报告', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'GET'],
            ['路径', '/api/study/report'],
            ['认证', '需要'],
            ['说明', '获取指定日期的学习报告'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '查询参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['参数', '类型', '必填', '说明'],
        [
            ['report_date', 'string', '否', '日期格式 YYYY-MM-DD，默认今天'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "date": "2026-07-03",
  "total_focus_minutes": 120,
  "distraction_count": 5,
  "sessions": [
    {
      "id": 1,
      "start_time": "2026-07-03T09:00:00",
      "end_time": "2026-07-03T10:00:00",
      "focus_duration_seconds": 3600,
      "distraction_count": 2
    }
  ]
}''')
    
    add_heading_styled(doc, '七、运行时设置模块', level=1)
    
    add_heading_styled(doc, '7.1 设置 API Key', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'POST'],
            ['路径', '/api/settings/apikey'],
            ['认证', '需要'],
            ['说明', '运行时动态设置 API Key（内存存储，不写入文件）'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '请求参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['字段', '类型', '必填', '说明'],
        [
            ['key', 'string', '是', 'Key 名称：DEEPSEEK_API_KEY 或 OPENAI_API_KEY'],
            ['value', 'string', '是', 'API Key 值（DeepSeek 需以 sk- 开头）'],
        ],
        col_widths=[3, 3, 2, 7]
    )
    
    add_normal_para(doc, '成功响应 (200)：', bold=True, indent=False)
    add_code_block(doc, '''{
  "status": "ok",
  "key": "DEEPSEEK_API_KEY",
  "configured": true
}''')
    
    add_heading_styled(doc, '7.2 查询 API Key 状态', level=2)
    add_table_simple(doc,
        ['项目', '内容'],
        [
            ['方法', 'GET'],
            ['路径', '/api/settings/apikey/{key_name}'],
            ['认证', '需要'],
            ['说明', '查询某个 API Key 是否已配置（不返回实际值）'],
        ],
        col_widths=[3, 12]
    )
    
    add_normal_para(doc, '路径参数：', bold=True, indent=False)
    add_table_simple(doc,
        ['参数', '类型', '说明'],
        [
            ['key_name', 'string', 'DEEPSEEK_API_KEY 或 OPENAI_API_KEY'],
        ],
        col_widths=[3, 3, 9]
    )
    
    output_path = r'D:\AI Study\.trae-cn\worktrees\AI Study\feat-yolov8-face-detection-KcMOL5\docs\04-前后端API接口定义表.docx'
    doc.save(output_path)
    print(f'文档4已生成: {output_path}')


if __name__ == '__main__':
    generate_doc1()
    generate_doc2()
    generate_doc3()
    generate_doc4()
    print('所有文档生成完成！')
